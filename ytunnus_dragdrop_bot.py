# protestibotti.py
# ProtestiBotti (Ultra HUD + Speed upgrades)
# Moodit:
#  1) PDF -> Y-tunnukset -> YTJ sähköpostit
#  2) Kauppalehti (Chrome debug 9222) -> Y-tunnukset -> YTJ sähköpostit
#  3) Clipboard -> (Y-tunnukset suoraan jos löytyy) muuten: yritysnimet -> YTJ: Y-tunnukset -> YTJ: sähköpostit
#
# Parannukset:
#  - Ultra modern HUD (header strip assets, centered layout, dark cards)
#  - Assets-kansio (toimii myös PyInstaller exe:ssä)
#  - Nopeutus: email etsitään ensin, klikataan "Näytä" vain jos pakko
#  - Cache: yt->email ja name->yt (säästää duplikaateissa)
#  - Tulokset aina uuteen docx: sahkopostit_001.docx, sahkopostit_002.docx, ...
#
# Riippuvuudet:
#   pip install selenium webdriver-manager PyPDF2 python-docx tkinterdnd2 pillow
#
# Build local:
#   pyinstaller --noconfirm --clean --onefile --windowed --name ProtestiBotti --add-data "assets;assets" protestibotti.py

import os
import re
import sys
import time
import threading
import subprocess
from difflib import SequenceMatcher

import tkinter as tk
from tkinter import ttk, messagebox, filedialog

import PyPDF2
from docx import Document

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException, WebDriverException
from webdriver_manager.chrome import ChromeDriverManager

# Drag & Drop (PDF)
try:
    from tkinterdnd2 import DND_FILES, TkinterDnD  # type: ignore
    HAS_DND = True
except Exception:
    HAS_DND = False

# Pillow for assets (jpg/png + scaling + glow)
try:
    from PIL import Image, ImageTk, ImageFilter, ImageOps  # type: ignore
    HAS_PIL = True
except Exception:
    HAS_PIL = False


# =========================
#   SPEED TUNING
# =========================
YTJ_PAGE_LOAD_TIMEOUT = 16
YTJ_RETRY_READS = 4
YTJ_RETRY_SLEEP = 0.10
YTJ_NAYTA_PASSES = 2
YTJ_PER_COMPANY_SLEEP = 0.02

KL_LOAD_MORE_WAIT = 1.0
KL_COMPANY_PAGE_TIMEOUT = 16
KL_AFTER_OPEN_SLEEP = 0.03

CLIP_YTJ_SEARCH_TIMEOUT = 8.0
CLIP_PER_NAME_SLEEP = 0.01

PARTIAL_SAVE_EVERY_NEW_EMAILS = 40


# =========================
#   REGEX
# =========================
YT_RE = re.compile(r"\b\d{7}-\d\b|\b\d{8}\b")
EMAIL_RE = re.compile(r"[A-Za-z0-9_.+-]+@[A-Za-z0-9-]+\.[A-Za-z0-9-.]+")
EMAIL_A_RE = re.compile(r"[A-Za-z0-9_.+-]+\s*\(a\)\s*[A-Za-z0-9-]+\.[A-Za-z0-9-.]+", re.I)

KAUPPALEHTI_URL = "https://www.kauppalehti.fi/yritykset/protestilista"
KAUPPALEHTI_MATCH = "kauppalehti.fi/yritykset/protestilista"
YTJ_COMPANY_URL = "https://tietopalvelu.ytj.fi/yritys/{}"

STRICT_FORMS_RE = re.compile(
    r"\b(oy|ab|ky|tmi|oyj|osakeyhtiö|kommandiittiyhtiö|toiminimi|as\.|ltd|llc|inc|gmbh)\b",
    re.IGNORECASE,
)


# =========================
#   PATHS / RESOURCES
# =========================
def resource_base_dir() -> str:
    # PyInstaller: sys._MEIPASS contains temp extraction dir
    if getattr(sys, "frozen", False) and hasattr(sys, "_MEIPASS"):
        return getattr(sys, "_MEIPASS")  # type: ignore
    return os.path.dirname(os.path.abspath(__file__))


def exe_dir() -> str:
    if getattr(sys, "frozen", False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))


ASSETS_DIR = os.path.join(resource_base_dir(), "assets")


def get_output_dir():
    base = exe_dir()
    try:
        p = os.path.join(base, "_write_test.tmp")
        with open(p, "w", encoding="utf-8") as f:
            f.write("ok")
        os.remove(p)
    except Exception:
        home = os.path.expanduser("~")
        docs = os.path.join(home, "Documents")
        base = os.path.join(docs, "ProtestiBotti")

    date_folder = time.strftime("%Y-%m-%d")
    out = os.path.join(base, date_folder)
    os.makedirs(out, exist_ok=True)
    return out


OUT_DIR = get_output_dir()
LOG_PATH = os.path.join(OUT_DIR, "log.txt")
EMAILS_TMP_PATH = os.path.join(OUT_DIR, "emails_tmp.txt")
PARTIAL_DOCX_PATH = os.path.join(OUT_DIR, "sahkopostit_partial.docx")


def reset_log():
    try:
        with open(LOG_PATH, "w", encoding="utf-8") as f:
            f.write("=== BOTTI KÄYNNISTETTY ===\n")
    except Exception:
        pass
    try:
        with open(EMAILS_TMP_PATH, "w", encoding="utf-8") as f:
            f.write("")
    except Exception:
        pass
    log_to_file(f"Output: {OUT_DIR}")
    log_to_file(f"Logi: {LOG_PATH}")
    log_to_file(f"Assets: {ASSETS_DIR}")


def log_to_file(msg: str):
    ts = time.strftime("%H:%M:%S")
    line = f"[{ts}] {msg}"
    try:
        with open(LOG_PATH, "a", encoding="utf-8") as f:
            f.write(line + "\n")
    except Exception:
        pass
    return line


# =========================
#   UNIQUE DOCX NAMES
# =========================
def next_indexed_docx(prefix: str, start_at: int = 1) -> str:
    i = start_at
    while True:
        name = f"{prefix}_{i:03d}.docx"
        path = os.path.join(OUT_DIR, name)
        if not os.path.exists(path):
            return path
        i += 1


def save_word_unique(lines, prefix: str):
    path = next_indexed_docx(prefix)
    doc = Document()
    for line in lines:
        if line:
            doc.add_paragraph(line)
    doc.save(path)
    return path


def save_word_to_path(lines, path):
    doc = Document()
    for line in lines:
        if line:
            doc.add_paragraph(line)
    doc.save(path)
    return path


def append_email_tmp(email: str):
    try:
        with open(EMAILS_TMP_PATH, "a", encoding="utf-8") as f:
            f.write(email.strip() + "\n")
    except Exception:
        pass


# =========================
#   TEXT UTIL
# =========================
def normalize_yt(yt: str):
    yt = (yt or "").strip().replace(" ", "")
    if re.fullmatch(r"\d{7}-\d", yt):
        return yt
    if re.fullmatch(r"\d{8}", yt):
        return yt[:7] + "-" + yt[7]
    return None


def pick_email_from_text(text: str) -> str:
    if not text:
        return ""
    m = EMAIL_RE.search(text)
    if m:
        return m.group(0).strip().replace(" ", "")
    m2 = EMAIL_A_RE.search(text)
    if m2:
        return m2.group(0).replace(" ", "").replace("(a)", "@")
    return ""


def split_lines(text: str):
    if not text:
        return []
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    return [ln.strip() for ln in text.split("\n") if ln.strip()]


def extract_yts_from_text(text: str):
    yts = set()
    for m in YT_RE.findall(text or ""):
        n = normalize_yt(m)
        if n:
            yts.add(n)
    return sorted(yts)


def extract_yt_from_text_anywhere(txt: str) -> str:
    if not txt:
        return ""
    for m in YT_RE.findall(txt):
        n = normalize_yt(m)
        if n:
            return n
    return ""


def extract_names_from_clipboard(text: str, strict: bool, max_names: int):
    lines = split_lines(text)
    out = []
    seen = set()

    bad_contains = [
        "näytä lisää", "protestilista", "kauppalehti", "kirjaudu", "tilaa", "tilaajille",
        "€", "eur", "summa", "viiväst", "päivä", "päivää", "päivämäärä",
        "y-tunnus", "y tunnus", "ytunnus", "osoite", "postinumero",
        "sähköposti", "puhelin", "www.", "http",
    ]

    for ln in lines:
        if len(out) >= max_names:
            break
        low = ln.lower()
        if YT_RE.search(ln):
            continue
        if any(b in low for b in bad_contains):
            continue
        if len(ln) < 3:
            continue
        digits = sum(ch.isdigit() for ch in ln)
        if digits >= 3:
            continue
        if not any(ch.isalpha() for ch in ln):
            continue

        name = re.sub(r"\s{2,}", " ", ln).strip()
        if len(name) > 90:
            continue
        if strict and not STRICT_FORMS_RE.search(name):
            continue

        key = name.lower()
        if key in seen:
            continue
        seen.add(key)
        out.append(name)

    return out


# =========================
#   PDF
# =========================
def extract_ytunnukset_from_pdf(pdf_path: str):
    yt_set = set()
    with open(pdf_path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            text = page.extract_text() or ""
            for m in YT_RE.findall(text):
                n = normalize_yt(m)
                if n:
                    yt_set.add(n)
    return sorted(yt_set)


# =========================
#   SELENIUM HELPERS
# =========================
def start_new_driver():
    options = webdriver.ChromeOptions()
    options.add_argument("--start-maximized")
    options.add_argument("--disable-gpu")
    options.add_argument("--disable-dev-shm-usage")
    driver_path = ChromeDriverManager().install()
    drv = webdriver.Chrome(service=Service(driver_path), options=options)
    drv.set_page_load_timeout(YTJ_PAGE_LOAD_TIMEOUT)
    return drv


def attach_to_existing_chrome():
    options = webdriver.ChromeOptions()
    options.add_experimental_option("debuggerAddress", "127.0.0.1:9222")
    driver_path = ChromeDriverManager().install()
    drv = webdriver.Chrome(service=Service(driver_path), options=options)
    drv.set_page_load_timeout(YTJ_PAGE_LOAD_TIMEOUT)
    return drv


def open_new_tab(driver, url="about:blank"):
    driver.execute_script("window.open(arguments[0], '_blank');", url)
    driver.switch_to.window(driver.window_handles[-1])


def safe_scroll_into_view(driver, elem):
    try:
        driver.execute_script("arguments[0].scrollIntoView({block:'center'});", elem)
    except Exception:
        pass


def safe_click(driver, elem) -> bool:
    try:
        safe_scroll_into_view(driver, elem)
        time.sleep(0.01)
        try:
            elem.click()
        except Exception:
            driver.execute_script("arguments[0].click();", elem)
        return True
    except Exception:
        return False


def try_accept_cookies(driver):
    texts = ["Hyväksy", "Hyväksy kaikki", "Salli kaikki", "Accept", "Accept all", "I agree", "OK", "Selvä"]
    for _ in range(2):
        found = False
        for e in driver.find_elements(By.XPATH, "//button|//a|//*[@role='button']"):
            try:
                t = (e.text or "").strip()
                if not t:
                    continue
                low = t.lower()
                if any(x.lower() in low for x in texts):
                    if e.is_displayed() and e.is_enabled():
                        safe_click(driver, e)
                        time.sleep(0.15)
                        found = True
                        break
            except Exception:
                continue
        if not found:
            break


# =========================
#   YTJ EMAIL (Speed upgrade)
# =========================
def wait_ytj_loaded(driver):
    wait = WebDriverWait(driver, YTJ_PAGE_LOAD_TIMEOUT)
    wait.until(EC.presence_of_element_located((By.TAG_NAME, "body")))


def extract_email_from_ytj(driver):
    # mailto first (fast)
    try:
        for a in driver.find_elements(By.TAG_NAME, "a"):
            href = (a.get_attribute("href") or "")
            if href.lower().startswith("mailto:"):
                return href.split(":", 1)[1].strip()
    except Exception:
        pass

    # row with label "Sähköposti"
    try:
        cells = driver.find_elements(By.XPATH, "//tr//*[self::td or self::th][contains(normalize-space(.), 'Sähköposti')]")
        for c in cells:
            try:
                tr = c.find_element(By.XPATH, "ancestor::tr[1]")
                email = pick_email_from_text(tr.text or "")
                if email:
                    return email
            except Exception:
                continue
    except Exception:
        pass

    # fallback whole body
    try:
        return pick_email_from_text(driver.find_element(By.TAG_NAME, "body").text or "")
    except Exception:
        return ""


def click_all_nayta_ytj(driver):
    for _ in range(YTJ_NAYTA_PASSES):
        clicked = False
        for b in driver.find_elements(By.TAG_NAME, "button"):
            try:
                if (b.text or "").strip().lower() == "näytä" and b.is_displayed() and b.is_enabled():
                    safe_click(driver, b)
                    clicked = True
                    time.sleep(0.06)
            except Exception:
                continue
        for a in driver.find_elements(By.TAG_NAME, "a"):
            try:
                if (a.text or "").strip().lower() == "näytä" and a.is_displayed():
                    safe_click(driver, a)
                    clicked = True
                    time.sleep(0.06)
            except Exception:
                continue
        if not clicked:
            break


def fetch_emails_from_ytj(driver, yt_list, status_cb, progress_cb, log_cb, stop_flag):
    emails = []
    seen_email = set()

    # cache: yt -> email
    yt_email_cache = {}

    progress_cb(0, max(1, len(yt_list)))
    new_since_partial = 0

    for i, yt in enumerate(yt_list, start=1):
        if stop_flag.is_set():
            status_cb("Pysäytetty.")
            break

        progress_cb(i - 1, len(yt_list))
        status_cb(f"YTJ email: {i}/{len(yt_list)} {yt}")

        if yt in yt_email_cache:
            em = yt_email_cache[yt]
            if em and em.lower() not in seen_email:
                seen_email.add(em.lower())
                emails.append(em)
                log_cb(em)
            continue

        try:
            driver.get(YTJ_COMPANY_URL.format(yt))
        except TimeoutException:
            pass

        try:
            wait_ytj_loaded(driver)
        except Exception:
            pass

        try_accept_cookies(driver)

        # SPEED: try extract BEFORE clicking Näytä
        email = ""
        for _ in range(max(1, YTJ_RETRY_READS - 1)):
            if stop_flag.is_set():
                break
            email = extract_email_from_ytj(driver)
            if email:
                break
            time.sleep(YTJ_RETRY_SLEEP)

        # if not found -> click Näytä -> retry
        if not email:
            click_all_nayta_ytj(driver)
            for _ in range(YTJ_RETRY_READS):
                if stop_flag.is_set():
                    break
                email = extract_email_from_ytj(driver)
                if email:
                    break
                time.sleep(YTJ_RETRY_SLEEP)

        yt_email_cache[yt] = email

        if email:
            k = email.lower()
            if k not in seen_email:
                seen_email.add(k)
                emails.append(email)
                append_email_tmp(email)
                log_cb(email)

                new_since_partial += 1
                if new_since_partial >= PARTIAL_SAVE_EVERY_NEW_EMAILS:
                    try:
                        save_word_to_path(emails, PARTIAL_DOCX_PATH)
                        log_cb(f"(partial) Tallennettu: {PARTIAL_DOCX_PATH}")
                    except Exception:
                        pass
                    new_since_partial = 0

        time.sleep(YTJ_PER_COMPANY_SLEEP)

    progress_cb(len(yt_list), max(1, len(yt_list)))
    try:
        if emails:
            save_word_to_path(emails, PARTIAL_DOCX_PATH)
    except Exception:
        pass
    return emails


# =========================
#   YTJ NAME -> YT (with cache)
# =========================
def ytj_open_search_home(driver):
    driver.get("https://tietopalvelu.ytj.fi/")
    try:
        wait_ytj_loaded(driver)
    except Exception:
        pass
    try_accept_cookies(driver)


def find_ytj_search_input(driver):
    xpaths = [
        "//input[@type='search']",
        "//input[contains(translate(@placeholder,'HAE','hae'),'hae')]",
        "//input[contains(translate(@aria-label,'HAE','hae'),'hae')]",
        "//input[contains(translate(@name,'HAE','hae'),'hae')]",
        "//form//input",
        "//input",
    ]
    for xp in xpaths:
        try:
            cands = driver.find_elements(By.XPATH, xp)
        except Exception:
            cands = []
        for inp in cands:
            try:
                if not inp.is_displayed() or not inp.is_enabled():
                    continue
                t = (inp.get_attribute("type") or "").lower()
                if t in ("hidden", "password", "checkbox", "radio", "submit", "button"):
                    continue
                return inp
            except Exception:
                continue
    return None


def score_result(name_query: str, card_text: str) -> float:
    txt = (card_text or "").strip()
    m = SequenceMatcher(None, (name_query or "").lower(), txt.lower()).ratio()
    score = m * 100.0
    if extract_yt_from_text_anywhere(txt):
        score += 20.0
    return score


def ytj_find_company_and_open_best(driver, name: str, stop_flag):
    ytj_open_search_home(driver)
    if stop_flag.is_set():
        return False

    inp = find_ytj_search_input(driver)
    if not inp:
        return False

    try:
        try:
            inp.clear()
        except Exception:
            pass
        inp.send_keys(name)
        inp.send_keys(u"\ue007")  # ENTER
    except Exception:
        return False

    best_link = None
    best_score = -1.0
    t0 = time.time()

    while time.time() - t0 < CLIP_YTJ_SEARCH_TIMEOUT:
        if stop_flag.is_set():
            return False

        try:
            try_accept_cookies(driver)
        except Exception:
            pass

        candidate_links = []
        for xp in ("//a[contains(@href,'/yritys/')]", "//a[contains(@href,'yritys')]"):
            try:
                candidate_links.extend(driver.find_elements(By.XPATH, xp))
            except Exception:
                pass

        checked = 0
        for a in candidate_links:
            if checked >= 30:
                break
            try:
                if not a.is_displayed():
                    continue
                href = (a.get_attribute("href") or "")
                if not href or "tietopalvelu.ytj.fi" not in href:
                    continue
                try:
                    card = a.find_element(By.XPATH, "ancestor::*[self::li or self::div or self::article][1]")
                    card_text = (card.text or "")
                except Exception:
                    card_text = (a.text or "")
                s = score_result(name, card_text)
                checked += 1
                if s > best_score:
                    best_score = s
                    best_link = a
            except Exception:
                continue

        if best_link:
            break
        time.sleep(0.15)

    if not best_link:
        return False

    try:
        safe_click(driver, best_link)
        try:
            wait_ytj_loaded(driver)
        except Exception:
            pass
        return True
    except Exception:
        return False


def ytj_name_to_yt(driver, name: str, stop_flag) -> str:
    ok = ytj_find_company_and_open_best(driver, name, stop_flag)
    if not ok:
        return ""
    try_accept_cookies(driver)
    try:
        body = driver.find_element(By.TAG_NAME, "body").text or ""
    except Exception:
        body = ""
    return extract_yt_from_text_anywhere(body)


# =========================
#   KAUPPALEHTI -> YT (robust: open company page in new tab)
# =========================
def focus_kauppalehti_tab(driver) -> bool:
    for handle in driver.window_handles:
        try:
            driver.switch_to.window(handle)
            url = (driver.current_url or "")
            if KAUPPALEHTI_MATCH in url:
                return True
        except Exception:
            continue
    return False


def page_looks_like_protestilista(driver) -> bool:
    try:
        rows = driver.find_elements(By.XPATH, "//table//tbody//tr")
        if rows and len(rows) >= 3:
            return True
    except Exception:
        pass
    try:
        for b in driver.find_elements(By.XPATH, "//button|//*[@role='button']"):
            if (b.text or "").strip().lower() == "näytä lisää":
                return True
    except Exception:
        pass
    return False


def page_looks_like_login_or_paywall(driver) -> bool:
    try:
        text = (driver.find_element(By.TAG_NAME, "body").text or "").lower()
        bad_words = ["kirjaudu", "tilaa", "tilaajille", "subscribe", "login", "digitilaus"]
        return any(w in text for w in bad_words)
    except Exception:
        return False


def ensure_protestilista_open_and_ready(driver, status_cb, log_cb, max_wait_seconds=900, stop_flag=None) -> bool:
    if focus_kauppalehti_tab(driver):
        status_cb("Löytyi protestilista-tab.")
    else:
        status_cb("Protestilista-tab ei löytynyt -> avaan protestilistan uuteen tabiin…")
        log_cb("AUTOFIX: opening protestilista in new tab")
        open_new_tab(driver, KAUPPALEHTI_URL)
        try:
            WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.TAG_NAME, "body")))
        except Exception:
            pass
        try_accept_cookies(driver)

    start = time.time()
    warned = False

    while True:
        if stop_flag and stop_flag.is_set():
            status_cb("Pysäytetty.")
            return False

        try:
            try_accept_cookies(driver)
        except Exception:
            pass

        if page_looks_like_protestilista(driver):
            status_cb("Protestilista valmis (taulukko näkyy).")
            return True

        if page_looks_like_login_or_paywall(driver) and not warned:
            warned = True
            status_cb("Kauppalehti vaatii kirjautumisen/tilaajanäkymän. Kirjaudu Chrome-bottiin (9222).")
            log_cb("Waiting for login…")
            try:
                messagebox.showinfo(
                    "Kirjaudu Kauppalehteen",
                    "Kirjaudu nyt Kauppalehteen AUKI OLEVAAN Chrome-bottiin (9222).\n"
                    "Kun protestilista näkyy, botti jatkaa."
                )
            except Exception:
                pass

        if time.time() - start > max_wait_seconds:
            status_cb("Aikakatkaisu: protestilista ei tullut näkyviin.")
            log_cb("ERROR: timeout waiting protestilista")
            return False

        time.sleep(2)


def click_nayta_lisaa(driver) -> bool:
    for b in driver.find_elements(By.XPATH, "//button|//*[@role='button']"):
        try:
            if not b.is_displayed() or not b.is_enabled():
                continue
            if (b.text or "").strip().lower() == "näytä lisää":
                safe_click(driver, b)
                return True
        except Exception:
            continue
    return False


def get_company_hrefs_from_visible_rows(driver):
    hrefs = []
    rows = driver.find_elements(By.XPATH, "//table//tbody//tr")
    for r in rows:
        try:
            if not r.is_displayed():
                continue
            links = r.find_elements(By.XPATH, ".//td[1]//a[contains(@href,'/yritykset/') and normalize-space(.)!='']")
            for a in links:
                href = (a.get_attribute("href") or "").strip()
                if href and "/yritykset/" in href:
                    hrefs.append(href)
        except Exception:
            continue
    out, seen = [], set()
    for h in hrefs:
        if h not in seen:
            seen.add(h)
            out.append(h)
    return out


def extract_yt_from_company_page_in_new_tab(driver, href: str, stop_flag):
    if stop_flag.is_set():
        return ""
    parent = driver.current_window_handle
    open_new_tab(driver, href)
    yt = ""
    try:
        t0 = time.time()
        while time.time() - t0 < KL_COMPANY_PAGE_TIMEOUT:
            if stop_flag.is_set():
                return ""
            try:
                if driver.find_elements(By.TAG_NAME, "body"):
                    break
            except Exception:
                pass
            time.sleep(0.08)

        try_accept_cookies(driver)
        time.sleep(KL_AFTER_OPEN_SLEEP)

        try:
            body = driver.find_element(By.TAG_NAME, "body").text or ""
        except Exception:
            body = ""
        yt = extract_yt_from_text_anywhere(body)
    finally:
        try:
            driver.close()
        except Exception:
            pass
        try:
            driver.switch_to.window(parent)
        except Exception:
            pass
    return yt


def collect_yts_from_kauppalehti(driver, status_cb, log_cb, stop_flag):
    try:
        WebDriverWait(driver, 20).until(EC.presence_of_element_located((By.TAG_NAME, "body")))
    except Exception:
        pass
    try_accept_cookies(driver)

    if not ensure_protestilista_open_and_ready(driver, status_cb, log_cb, stop_flag=stop_flag):
        return []

    collected = set()
    seen_hrefs = set()

    while True:
        if stop_flag.is_set():
            status_cb("Pysäytetty.")
            break

        hrefs = get_company_hrefs_from_visible_rows(driver)
        if not hrefs:
            status_cb("Kauppalehti: en löydä yrityslinkkejä. Vinkki: Clipboard-moodi.")
            break

        new_hrefs = [h for h in hrefs if h not in seen_hrefs]
        status_cb(f"Kauppalehti: linkkejä {len(hrefs)} | uudet {len(new_hrefs)} | YT {len(collected)}")

        got = 0
        for href in new_hrefs:
            if stop_flag.is_set():
                break
            seen_hrefs.add(href)
            yt = extract_yt_from_company_page_in_new_tab(driver, href, stop_flag)
            if yt and yt not in collected:
                collected.add(yt)
                got += 1
                log_cb(f"+ {yt} (yht {len(collected)})")

        if stop_flag.is_set():
            status_cb("Pysäytetty.")
            break

        if click_nayta_lisaa(driver):
            status_cb("Kauppalehti: Näytä lisää…")
            time.sleep(KL_LOAD_MORE_WAIT)
            try:
                driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
            except Exception:
                pass
            time.sleep(0.2)
            continue

        if got == 0 and not new_hrefs:
            status_cb("Kauppalehti: valmis.")
            break

    return sorted(collected)


# =========================
#   CHROME BOT (9222)
# =========================
def build_chrome_bot_args():
    base = exe_dir()
    profile_dir = os.path.join(base, "chrome_bot_profile")
    os.makedirs(profile_dir, exist_ok=True)

    chrome_paths = [
        r"C:\Program Files\Google\Chrome\Application\chrome.exe",
        r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
    ]
    chrome_exe = None
    for p in chrome_paths:
        if os.path.exists(p):
            chrome_exe = p
            break
    if chrome_exe is None:
        chrome_exe = "chrome"

    return [
        chrome_exe,
        "--remote-debugging-port=9222",
        f"--user-data-dir={profile_dir}",
    ]


def launch_chrome_bot():
    try:
        args = build_chrome_bot_args()
        subprocess.Popen(args, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        return True
    except Exception:
        return False


# =========================
#   ULTRA HUD UI
# =========================
BaseTk = TkinterDnD.Tk if HAS_DND else tk.Tk


class ModernCard(tk.Frame):
    def __init__(self, parent, bg="#0b0b0f", border="#1c1c22", *args, **kwargs):
        super().__init__(parent, bg=bg, highlightthickness=1, highlightbackground=border, *args, **kwargs)


class App(BaseTk):
    def __init__(self):
        super().__init__()
        reset_log()

        self.stop_flag = threading.Event()

        # Colors
        self.BG = "#050507"
        self.CARD = "#0b0b0f"
        self.CARD2 = "#0e0e14"
        self.BORDER = "#1f1f27"
        self.TEXT = "#f3f3f5"
        self.MUTED = "#b7b7c2"
        self.RED = "#cf1030"
        self.RED_H = "#ff2448"
        self.GREEN = "#19c37d"

        self.title("ProtestiBotti")
        self.geometry("1040x920")
        self.configure(bg=self.BG)
        self.bind_all("<Escape>", lambda e: self.request_stop())

        self._img_refs = {}
        self._header_labels = []

        # Root layout
        container = tk.Frame(self, bg=self.BG)
        container.pack(fill="both", expand=True)

        # Header strip
        header = ModernCard(container, bg=self.CARD, border=self.BORDER)
        header.pack(fill="x", padx=14, pady=(14, 10))

        header_grid = tk.Frame(header, bg=self.CARD)
        header_grid.pack(fill="x", padx=12, pady=10)

        # Title row + status pill
        title_row = tk.Frame(header_grid, bg=self.CARD)
        title_row.grid(row=0, column=0, sticky="ew")
        header_grid.grid_columnconfigure(0, weight=1)

        tk.Label(
            title_row, text="ProtestiBotti", bg=self.CARD, fg=self.TEXT,
            font=("Segoe UI", 22, "bold")
        ).pack(side="left")

        self.pill = tk.Label(
            title_row, text="VALMIINA", bg="#0f241c", fg=self.GREEN,
            font=("Segoe UI", 10, "bold"), padx=10, pady=6
        )
        self.pill.pack(side="right")

        tk.Label(
            header_grid,
            text="PDF / Kauppalehti / Clipboard → YTJ → sähköpostit (Näytä klikataan automaattisesti)",
            bg=self.CARD, fg=self.MUTED, font=("Segoe UI", 10)
        ).grid(row=1, column=0, sticky="w", pady=(4, 8))

        # Image strip (top row, no overlap)
        img_strip = tk.Frame(header_grid, bg=self.CARD)
        img_strip.grid(row=2, column=0, sticky="ew")
        header_grid.grid_rowconfigure(2, weight=0)

        # Create placeholders for images (we'll resize on window changes)
        self._create_header_image(img_strip, "h1.jpg")
        self._create_header_image(img_strip, "h.jpg")
        self._create_header_image(img_strip, "polis.jpg")
        self._create_header_image(img_strip, "uo.png")
        self._create_header_image(img_strip, "vero.png")

        img_strip.bind("<Configure>", self._resize_header_images)

        # Buttons card
        actions = ModernCard(container, bg=self.CARD, border=self.BORDER)
        actions.pack(fill="x", padx=14, pady=(0, 10))

        btn_row = tk.Frame(actions, bg=self.CARD)
        btn_row.pack(fill="x", padx=12, pady=12)

        self._mk_btn(btn_row, "Avaa Chrome-botti (9222)", self.open_chrome_bot).pack(side="left", padx=6)
        self._mk_btn(btn_row, "Kauppalehti → YTJ", self.start_kauppalehti_mode).pack(side="left", padx=6)
        self._mk_btn(btn_row, "PDF → YTJ", self.start_pdf_mode).pack(side="left", padx=6)
        self._mk_btn(btn_row, "Pysäytä", self.request_stop, danger=True).pack(side="right", padx=6)

        # Status + progress
        status_card = ModernCard(container, bg=self.CARD, border=self.BORDER)
        status_card.pack(fill="x", padx=14, pady=(0, 10))

        status_row = tk.Frame(status_card, bg=self.CARD)
        status_row.pack(fill="x", padx=12, pady=10)

        self.status = tk.Label(status_row, text="Valmiina.", bg=self.CARD, fg=self.TEXT, font=("Segoe UI", 11))
        self.status.pack(side="left")

        self.progress = ttk.Progressbar(status_card, orient="horizontal", mode="determinate", length=980)
        self.progress.pack(fill="x", padx=12, pady=(0, 12))

        # PDF drop
        pdf_card = ModernCard(container, bg=self.CARD, border=self.BORDER)
        pdf_card.pack(fill="x", padx=14, pady=(0, 10))

        self.drop_var = tk.StringVar(value="PDF: Pudota tähän (tai paina PDF → YTJ ja valitse tiedosto)")
        drop = tk.Label(
            pdf_card, textvariable=self.drop_var, bg=self.CARD2, fg=self.TEXT,
            font=("Segoe UI", 10), padx=12, pady=10,
            highlightthickness=1, highlightbackground=self.BORDER
        )
        drop.pack(fill="x", padx=12, pady=12)

        if HAS_DND:
            drop.drop_target_register(DND_FILES)
            drop.dnd_bind("<<Drop>>", self._on_drop_pdf)

        # Clipboard card
        clip_card = ModernCard(container, bg=self.CARD, border=self.BORDER)
        clip_card.pack(fill="x", padx=14, pady=(0, 10))

        top = tk.Frame(clip_card, bg=self.CARD)
        top.pack(fill="x", padx=12, pady=(12, 6))

        self.strict_var = tk.BooleanVar(value=True)
        tk.Checkbutton(
            top,
            text="Tiukka parsinta (vain Oy/Ab/Ky/Tmi/...)",
            variable=self.strict_var,
            bg=self.CARD, fg=self.TEXT,
            selectcolor=self.CARD2,
            activebackground=self.CARD,
            activeforeground=self.TEXT
        ).pack(side="left")

        tk.Label(top, text="Max nimeä:", bg=self.CARD, fg=self.TEXT, font=("Segoe UI", 10)).pack(side="left", padx=(16, 6))
        self.max_names_var = tk.IntVar(value=400)
        tk.Spinbox(
            top, from_=10, to=5000, textvariable=self.max_names_var, width=7,
            bg=self.CARD2, fg=self.TEXT, insertbackground=self.TEXT
        ).pack(side="left")

        tk.Label(clip_card, text="Clipboard (Ctrl+V tähän):", bg=self.CARD, fg=self.MUTED, font=("Segoe UI", 10)).pack(
            anchor="w", padx=12, pady=(6, 6)
        )

        self.clip_text = tk.Text(
            clip_card, height=8, wrap="word",
            bg=self.CARD2, fg=self.TEXT, insertbackground=self.TEXT,
            highlightthickness=1, highlightbackground=self.BORDER
        )
        self.clip_text.pack(fill="x", padx=12, pady=(0, 10))

        btm = tk.Frame(clip_card, bg=self.CARD)
        btm.pack(fill="x", padx=12, pady=(0, 12))
        self._mk_btn(btm, "Clipboard → (Nimet/YT → Email)", self.start_clipboard_mode).pack(side="left")

        # Log card
        log_card = ModernCard(container, bg=self.CARD, border=self.BORDER)
        log_card.pack(fill="both", expand=True, padx=14, pady=(0, 14))

        tk.Label(log_card, text="Live-logi (uusimmat alimmaisena):", bg=self.CARD, fg=self.MUTED, font=("Segoe UI", 10)).pack(
            anchor="w", padx=12, pady=(12, 6)
        )

        body = tk.Frame(log_card, bg=self.CARD)
        body.pack(fill="both", expand=True, padx=12, pady=(0, 12))

        self.listbox = tk.Listbox(
            body, height=14, bg=self.CARD2, fg=self.TEXT,
            highlightthickness=1, highlightbackground=self.BORDER,
            selectbackground=self.RED
        )
        self.listbox.pack(side="left", fill="both", expand=True)

        sb = ttk.Scrollbar(body, orient="vertical", command=self.listbox.yview)
        sb.pack(side="right", fill="y")
        self.listbox.configure(yscrollcommand=sb.set)

        self.ui_log(f"Tallennus: {OUT_DIR}")
        if not HAS_PIL:
            self.ui_log("HUOM: Pillow puuttuu -> .jpg + glow ei toimi. Asenna: pip install pillow")

        # Force initial header resize after window shows
        self.after(200, lambda: self._resize_header_images(None))

    # ---------- Modern widgets ----------
    def _mk_btn(self, parent, text, cmd, danger=False):
        bg = self.RED if not danger else "#7a0f1b"
        hover = self.RED_H if not danger else "#a81627"
        b = tk.Button(
            parent, text=text, command=cmd,
            bg=bg, fg=self.TEXT,
            activebackground=hover, activeforeground=self.TEXT,
            relief="flat", padx=14, pady=10,
            font=("Segoe UI", 10, "bold"),
        )
        b.bind("<Enter>", lambda e: b.configure(bg=hover))
        b.bind("<Leave>", lambda e: b.configure(bg=bg))
        return b

    def _create_header_image(self, parent, filename: str):
        lbl = tk.Label(parent, bg=self.CARD, fg=self.MUTED, font=("Segoe UI", 9))
        lbl.pack(side="left", padx=8, pady=2)
        self._header_labels.append((lbl, filename))

    def _load_asset(self, filename: str):
        path = os.path.join(ASSETS_DIR, filename)
        return path if os.path.exists(path) else None

    def _render_glow(self, img_rgba, glow_alpha=160):
        alpha = img_rgba.split()[-1]
        glow = Image.new("RGBA", img_rgba.size, (255, 0, 40, glow_alpha))
        glow.putalpha(alpha)
        glow = glow.filter(ImageFilter.GaussianBlur(radius=6))
        return Image.alpha_composite(glow, img_rgba)

    def _resize_header_images(self, event):
        if not HAS_PIL:
            # show missing pillow text once
            for lbl, fn in self._header_labels:
                lbl.configure(text=f"[{fn}]")
            return

        # Calculate width per image dynamically
        # Keep all visible, no overlap: allocate strip width / N
        try:
            strip = self._header_labels[0][0].master  # parent frame
            total_w = strip.winfo_width()
        except Exception:
            total_w = 900

        n = max(1, len(self._header_labels))
        # leave padding: 16px per image approx
        per = max(90, int((total_w - (n * 16)) / n))
        target_h = 78

        for lbl, fn in self._header_labels:
            p = self._load_asset(fn)
            if not p:
                lbl.configure(text=f"[puuttuu: {fn}]", image="")
                continue
            try:
                img = Image.open(p)
                img = ImageOps.contain(img, (per, target_h))
                if img.mode != "RGBA":
                    img = img.convert("RGBA")

                # If jpg, avoid rectangle glow: create alpha mask from luminance
                if fn.lower().endswith((".jpg", ".jpeg")):
                    gray = img.convert("L")
                    mask = gray.point(lambda px: 255 if px < 245 else 0)
                    img.putalpha(mask)

                out = self._render_glow(img, glow_alpha=150)
                photo = ImageTk.PhotoImage(out)
                lbl.configure(image=photo, text="")
                self._img_refs[fn] = photo
            except Exception as e:
                lbl.configure(text=f"[ei voi näyttää: {fn}]", image="")
                self.ui_log(f"KUVA ERROR {fn}: {e}")

    # ---------- UI helpers ----------
    def request_stop(self):
        self.stop_flag.set()
        self.ui_log("STOP: käyttäjä pyysi pysäytystä.")
        self.set_pill("PYSÄYTETÄÄN", warn=True)
        self.status.config(text="Pysäytetään…")

    def clear_stop(self):
        self.stop_flag.clear()
        self.set_pill("VALMIINA", ok=True)

    def set_pill(self, text, ok=False, warn=False):
        if ok:
            self.pill.configure(text=text, bg="#0f241c", fg=self.GREEN)
        elif warn:
            self.pill.configure(text=text, bg="#2a1215", fg=self.RED_H)
        else:
            self.pill.configure(text=text, bg="#14141a", fg=self.MUTED)

    def ui_log(self, msg):
        line = log_to_file(msg)
        try:
            self.listbox.insert(tk.END, line)
            self.listbox.yview_moveto(1.0)
            self.update_idletasks()
        except Exception:
            pass

    def set_status(self, s):
        self.status.config(text=s)
        self.update_idletasks()
        self.ui_log(s)

    def set_progress(self, value, maximum):
        self.progress["maximum"] = maximum
        self.progress["value"] = value
        self.update_idletasks()

    # ---------- actions ----------
    def open_chrome_bot(self):
        ok = launch_chrome_bot()
        if ok:
            self.set_pill("CHROME AUKI", ok=True)
            self.set_status("Chrome-botti avattu (9222). Kirjaudu Kauppalehteen siinä ikkunassa.")
        else:
            self.set_pill("VIRHE", warn=True)
            self.set_status("Chrome-botin avaus epäonnistui.")
            messagebox.showerror("Virhe", "Chrome-botin avaus epäonnistui. Tarkista Chrome-asennus.")

    def _on_drop_pdf(self, event):
        path = (event.data or "").strip()
        if path.startswith("{") and path.endswith("}"):
            path = path[1:-1]
        if path.lower().endswith(".pdf") and os.path.exists(path):
            self.drop_var.set(f"PDF valittu: {path}")
            self.clear_stop()
            threading.Thread(target=self.run_pdf_mode, args=(path,), daemon=True).start()
        else:
            messagebox.showwarning("Ei PDF", "Pudotettu tiedosto ei ollut .pdf")

    def start_pdf_mode(self):
        self.clear_stop()
        path = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf")])
        if path:
            self.drop_var.set(f"PDF valittu: {path}")
            threading.Thread(target=self.run_pdf_mode, args=(path,), daemon=True).start()

    def run_pdf_mode(self, pdf_path):
        driver = None
        try:
            self.set_pill("AJOSSA", warn=False)
            self.set_status("Luetaan PDF ja kerätään Y-tunnukset…")
            yt_list = extract_ytunnukset_from_pdf(pdf_path)

            if self.stop_flag.is_set():
                self.set_status("Pysäytetty.")
                return

            if not yt_list:
                self.set_status("Ei löytynyt Y-tunnuksia PDF:stä.")
                messagebox.showwarning("Ei löytynyt", "PDF:stä ei löytynyt yhtään Y-tunnusta.")
                return

            self.set_status("Käynnistetään Chrome ja haetaan sähköpostit YTJ:stä…")
            driver = start_new_driver()

            emails = fetch_emails_from_ytj(driver, yt_list, self.set_status, self.set_progress, self.ui_log, self.stop_flag)

            if self.stop_flag.is_set():
                self.set_status("Pysäytetty.")
                return

            em_path = save_word_unique(emails, "sahkopostit")
            self.ui_log(f"Tallennettu: {em_path}")

            self.set_pill("VALMIS", ok=True)
            self.set_status("Valmis!")
            messagebox.showinfo("Valmis", f"Valmis!\n\nKansio:\n{OUT_DIR}\n\nSähköposteja: {len(emails)}")

        except Exception as e:
            self.set_pill("VIRHE", warn=True)
            self.ui_log(f"VIRHE: {e}")
            self.set_status("Virhe. Katso log.txt")
            messagebox.showerror("Virhe", f"Tuli virhe.\nKatso log.txt:\n{LOG_PATH}\n\n{e}")
        finally:
            if driver:
                try:
                    driver.quit()
                except Exception:
                    pass

    def start_kauppalehti_mode(self):
        self.clear_stop()
        threading.Thread(target=self.run_kauppalehti_mode, daemon=True).start()

    def run_kauppalehti_mode(self):
        driver = None
        try:
            self.set_pill("AJOSSA", warn=False)
            self.set_status("Liitytään Chrome-bottiin (9222)…")
            driver = attach_to_existing_chrome()

            self.set_status("Kauppalehti: kerätään Y-tunnukset…")
            yt_list = collect_yts_from_kauppalehti(driver, self.set_status, self.ui_log, self.stop_flag)

            if self.stop_flag.is_set():
                self.set_status("Pysäytetty.")
                return

            if not yt_list:
                self.set_status("Ei löytynyt Y-tunnuksia. Vinkki: Clipboard-moodi.")
                messagebox.showwarning("Ei löytynyt", "Y-tunnuksia ei saatu. Kokeile Clipboard-moodia (Ctrl+A Ctrl+C -> Ctrl+V).")
                return

            self.set_status("Avataan YTJ uuteen tabiin…")
            open_new_tab(driver, "about:blank")

            self.set_status("YTJ: haetaan sähköpostit…")
            emails = fetch_emails_from_ytj(driver, yt_list, self.set_status, self.set_progress, self.ui_log, self.stop_flag)

            if self.stop_flag.is_set():
                self.set_status("Pysäytetty.")
                return

            em_path = save_word_unique(emails, "sahkopostit")
            self.ui_log(f"Tallennettu: {em_path}")

            self.set_pill("VALMIS", ok=True)
            self.set_status("Valmis!")
            messagebox.showinfo("Valmis", f"Valmis!\n\nKansio:\n{OUT_DIR}\n\nSähköposteja: {len(emails)}")

        except WebDriverException as e:
            self.set_pill("VIRHE", warn=True)
            self.ui_log(f"SELENIUM VIRHE: {e}")
            self.set_status("Virhe. Katso log.txt")
            messagebox.showerror("Virhe", f"Selenium/Chrome virhe.\nKatso log.txt:\n{LOG_PATH}\n\n{e}")
        except Exception as e:
            self.set_pill("VIRHE", warn=True)
            self.ui_log(f"VIRHE: {e}")
            self.set_status("Virhe. Katso log.txt")
            messagebox.showerror("Virhe", f"Tuli virhe.\nKatso log.txt:\n{LOG_PATH}\n\n{e}")
        finally:
            # do not quit attached chrome
            pass

    def start_clipboard_mode(self):
        self.clear_stop()
        text = self.clip_text.get("1.0", tk.END).strip()
        if not text:
            messagebox.showwarning("Tyhjä", "Liitä ensin teksti (Ctrl+V) kenttään.")
            return
        threading.Thread(target=self.run_clipboard_mode, args=(text,), daemon=True).start()

    def run_clipboard_mode(self, text: str):
        driver = None
        try:
            self.set_pill("AJOSSA", warn=False)

            # 1) If clipboard already contains YTs -> go straight to email
            yt_list_direct = extract_yts_from_text(text)
            if yt_list_direct:
                self.set_status(f"Clipboard: löytyi {len(yt_list_direct)} Y-tunnusta → haetaan sähköpostit…")
                driver = start_new_driver()
                emails = fetch_emails_from_ytj(driver, yt_list_direct, self.set_status, self.set_progress, self.ui_log, self.stop_flag)

                if self.stop_flag.is_set():
                    self.set_status("Pysäytetty.")
                    return

                em_path = save_word_unique(emails, "sahkopostit")
                self.ui_log(f"Tallennettu: {em_path}")

                self.set_pill("VALMIS", ok=True)
                self.set_status("Valmis!")
                messagebox.showinfo("Valmis", f"Valmis!\n\nKansio:\n{OUT_DIR}\n\nSähköposteja: {len(emails)}")
                return

            # 2) Otherwise: names -> yts -> emails
            strict = bool(self.strict_var.get())
            max_names = int(self.max_names_var.get() or 400)
            names = extract_names_from_clipboard(text, strict=strict, max_names=max_names)

            if not names:
                self.set_status("Clipboard: en löytänyt yritysnimiä.")
                messagebox.showwarning("Ei löytynyt", "Tekstistä ei saatu irti yritysnimiä.\nKokeile ottaa Tiukka parsinta pois.")
                return

            self.set_status(f"Clipboard: haetaan Y-tunnukset YTJ:stä nimillä… (nimiä {len(names)})")
            driver = start_new_driver()

            name_to_yt_cache = {}
            yts = []
            seen_yts = set()

            self.set_progress(0, max(1, len(names)))

            for i, name in enumerate(names, start=1):
                if self.stop_flag.is_set():
                    self.set_status("Pysäytetty.")
                    return

                self.set_progress(i - 1, len(names))
                self.set_status(f"YTJ Y-tunnus: {i}/{len(names)}  {name}")

                if name in name_to_yt_cache:
                    yt = name_to_yt_cache[name]
                else:
                    yt = ytj_name_to_yt(driver, name, self.stop_flag)
                    name_to_yt_cache[name] = yt

                if yt and yt not in seen_yts:
                    seen_yts.add(yt)
                    yts.append(yt)
                    self.ui_log(f"+YT {yt}  ({len(yts)})")
                elif not yt:
                    self.ui_log(f"YT NOT FOUND: {name}")

                time.sleep(CLIP_PER_NAME_SLEEP)

            self.set_progress(len(names), max(1, len(names)))

            if not yts:
                self.set_status("Ei löytynyt yhtään Y-tunnusta nimillä.")
                messagebox.showwarning("Ei löytynyt", "YTJ-nimihaulla ei saatu yhtään Y-tunnusta.\nKokeile: Tiukka parsinta pois.")
                return

            self.set_status(f"YTJ: haetaan sähköpostit {len(yts)} Y-tunnuksella…")
            emails = fetch_emails_from_ytj(driver, yts, self.set_status, self.set_progress, self.ui_log, self.stop_flag)

            if self.stop_flag.is_set():
                self.set_status("Pysäytetty.")
                return

            em_path = save_word_unique(emails, "sahkopostit")
            self.ui_log(f"Tallennettu: {em_path}")

            self.set_pill("VALMIS", ok=True)
            self.set_status("Valmis!")
            messagebox.showinfo(
                "Valmis",
                f"Valmis!\n\nKansio:\n{OUT_DIR}\n\nNimiä: {len(names)}\nY-tunnuksia: {len(yts)}\nSähköposteja: {len(emails)}"
            )

        except Exception as e:
            self.set_pill("VIRHE", warn=True)
            self.ui_log(f"VIRHE: {e}")
            self.set_status("Virhe. Katso log.txt")
            messagebox.showerror("Virhe", f"Tuli virhe.\nKatso log.txt:\n{LOG_PATH}\n\n{e}")
        finally:
            if driver:
                try:
                    driver.quit()
                except Exception:
                    pass


if __name__ == "__main__":
    App().mainloop()
